import os
import re
import json
import glob
import time
import pandas as pd
from docx import Document
from PyPDF2 import PdfReader
from groq import Groq
from dotenv import load_dotenv

# Load ENV from the local app directory
dotenv_path = r"C:\Users\Ali Musthaq\Downloads\WL OPS - for antigravity\wl-local-app\.env"
load_dotenv(dotenv_path)

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
if not GROQ_API_KEY:
    print("WARNING: GROQ_API_KEY not found in .env. Will try to extract from text directly.")
    try:
        with open(dotenv_path, "r") as f:
            for line in f:
                if "GROQ_API_KEY=" in line:
                    GROQ_API_KEY = line.split("=")[1].strip()
                    break
    except Exception as e:
        print(f"Error parsing .env manually: {e}")

if not GROQ_API_KEY:
    raise ValueError("GROQ_API_KEY is absolutely required!")

client = Groq(api_key=GROQ_API_KEY)

BASE_DIR = r"C:\Users\Ali Musthaq\Downloads\WL OPS - for antigravity\wellland activities"

def should_parse_file(file_path):
    name = os.path.basename(file_path).lower()
    path = file_path.lower()
    
    # Exclude SOP, Forms, Performance Appraisal, and Payslip folders
    if any(k in path for k in ["sop pack", "forms and templates", "performance appraisal", "payslip", "personal_info"]):
        return False
        
    # Exclude temp Excel/Word lock files
    if name.startswith("~$"):
        return False
        
    # Exclude duplicates, delivery notes, and giant consolidated files
    if any(k in name for k in ["deliverynote", "delivery note", "all inv", "all_inv", "signed", "compressed", "duplicate"]):
        return False
        
    # Skip draft or revision duplicates (e.g. file (1).pdf, file (2).pdf)
    if re.search(r"\(\d+\)", name):
        return False
        
    # Include files with relevant keywords
    keywords = [
        "expense", "sales", "invoice", "quotation", "quote", "log", "issue", 
        "todo", "outstanding", "shipment", "register", "tracker", "item", 
        "list", "status", "pr-", "po-", "inv-", "lot-", "delivery", "receipt", 
        "details", "operators", "vehicles", "staff", "asset", "master", "operations"
    ]
    if any(k in name for k in keywords):
        return True
        
    return False

def extract_text_from_docx(file_path):
    print(f"Extracting DOCX: {file_path}")
    doc = Document(file_path)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            fullText.append(" | ".join(row_text))
    return "\n".join(fullText)

def extract_text_from_pdf(file_path):
    print(f"Extracting PDF: {file_path}")
    reader = PdfReader(file_path)
    fullText = []
    # limit to first 3 pages to avoid huge contexts
    for page in reader.pages[:3]:
        text = page.extract_text()
        if text:
            fullText.append(text)
    return "\n".join(fullText)

def extract_tables_from_xlsx(file_path):
    print(f"Extracting XLSX: {file_path}")
    xl = pd.ExcelFile(file_path)
    output = []
    for sheet_name in xl.sheet_names:
        # Skip irrelevant sheets to save huge tokens
        if any(k in sheet_name.lower() for k in ["staff", "operator", "location", "contact", "phone", "template", "sop"]):
            print(f"  Skipping sheet: {sheet_name}")
            continue
        df = xl.parse(sheet_name)
        # Drop completely empty rows and columns
        df = df.dropna(how='all').dropna(axis=1, how='all')
        if df.empty:
            continue
        output.append(f"### Sheet: {sheet_name} ###")
        # Keep top 50 rows to save space, represent as CSV
        csv_str = df.head(50).to_csv(index=False)
        output.append(csv_str)
    return "\n\n".join(output)

def main():
    # Find all main files recursively
    docx_files = glob.glob(os.path.join(BASE_DIR, "**", "*.docx"), recursive=True)
    xlsx_files = glob.glob(os.path.join(BASE_DIR, "**", "*.xlsx"), recursive=True)
    pdf_files = glob.glob(os.path.join(BASE_DIR, "**", "*.pdf"), recursive=True)
    
    # We will aggregate all extracted information
    all_extracted_text = []
    
    # Process DOCX
    for docx in docx_files:
        name = os.path.basename(docx)
        if not should_parse_file(docx):
            continue
        if os.path.getsize(docx) > 2300000:
            print(f"Skipping giant DOCX: {name}")
            continue
        try:
            text = extract_text_from_docx(docx)
            all_extracted_text.append(f"--- File: {name} ---\n{text}")
        except Exception as e:
            print(f"Error reading docx {name}: {e}")
        
    # Process PDF
    for pdf in pdf_files:
        name = os.path.basename(pdf)
        if not should_parse_file(pdf):
            continue
        if any(k in name for k in ["EXCAVATOR", "KOBELCO", "Transformation", "Migration", "Blueprint", "Readiness", "rates deve"]):
            print(f"Skipping manual/blueprint PDF: {name}")
            continue
        if os.path.getsize(pdf) > 2300000:
            print(f"Skipping giant PDF: {name}")
            continue
        try:
            text = extract_text_from_pdf(pdf)
            all_extracted_text.append(f"--- File: {name} ---\n{text}")
        except Exception as e:
            print(f"Error reading pdf {name}: {e}")
        
    # Process XLSX
    for xlsx in xlsx_files:
        name = os.path.basename(xlsx)
        if not should_parse_file(xlsx):
            continue
        if os.path.getsize(xlsx) > 2300000:
            print(f"Skipping giant XLSX: {name}")
            continue
        try:
            text = extract_tables_from_xlsx(xlsx)
            all_extracted_text.append(f"--- File: {name} ---\n{text}")
        except Exception as e:
            print(f"Error reading xlsx {name}: {e}")
            
    combined_raw_text = "\n\n".join(all_extracted_text)
    
    # Save combined text for reference
    with open("combined_raw_extracted.txt", "w", encoding="utf-8") as f:
        f.write(combined_raw_text)
    print(f"Saved combined raw text to combined_raw_extracted.txt (Total chars: {len(combined_raw_text)})")
    
    # Now use Groq to parse into structured chunks in a single pass
    combined_prompt = """
    You are an expert data extraction assistant. You are given raw extracted text/tables from business documents.
    Extract the following three entities:
    1. Suppliers/Vendors: All companies or individuals selling goods or services.
    2. Inventory & Procurement: Items purchased, requested, or stored.
    3. Issue Reports: Breakdowns, maintenance requests, and required parts to rectify them.
    
    Format the output strictly as a JSON object with this exact structure:
    {
      "suppliers": [
        {
          "Supplier": "Name of supplier",
          "Category": "MRO / Electrical / Hydraulic / Rental etc",
          "Contact": "Contact number or email if available",
          "Remarks": "What they do or what was purchased from them"
        }
      ],
      "inventory": [
        {
          "Item Name": "Name of part/item",
          "Part Number": "Part number",
          "Category": "MRO / Engine Parts / Hydraulic Filters / Tyres etc",
          "Supplier": "Supplier name",
          "Unit Price": 0,
          "Quantity": 1,
          "Location": "Thilafushi / Tug-01 / Male"
        }
      ],
      "issueReports": [
        {
          "title": "Short title of issue",
          "description": "Full description of issue",
          "category": "Mechanical / Electrical / Hydraulic / Structural / Tyres / Undercarriage",
          "priority": "CRITICAL / HIGH / MEDIUM / LOW",
          "status": "Reported / Resolved / Under Review",
          "assetId": "Machine Asset ID if matching (e.g. EX-01, SK-02, CR-01). Check if asset name or ID matches. If no asset, leave empty.",
          "location": "Thilafushi Site / Tug-01 / Male etc",
          "reportedBy": "Who reported (e.g. Janaka, Tharanga, Mushthaq)",
          "parts": [
            {
              "name": "Part name required",
              "partNumber": "Part number or OEM number",
              "qty": 1,
              "estimatedCost": 0,
              "preferredSupplier": "Name of supplier if mentioned"
            }
          ]
        }
      ]
    }
    
    DO NOT wrap in any markdown backticks except the JSON block. Return ONLY valid JSON.
    """

    results_suppliers = []
    results_inventory = []
    results_issues = []
    
    chunk_size = 7000  # Smaller chunks to fit well within Groq's 6,000 TPM limit
    for i in range(0, len(combined_raw_text), chunk_size):
        chunk = combined_raw_text[i:i+chunk_size]
        payload = f"{combined_prompt}\n\nRaw Text Chunk:\n{chunk}"
        try:
            chat_completion = client.chat.completions.create(
                messages=[{"role": "user", "content": payload}],
                model="llama-3.1-8b-instant",  # Use Llama 3.1 8B with high TPD limit
                response_format={"type": "json_object"},
                temperature=0.1
            )
            res = json.loads(chat_completion.choices[0].message.content)
            
            # Merge lists
            if "suppliers" in res and isinstance(res["suppliers"], list):
                results_suppliers.extend(res["suppliers"])
            if "inventory" in res and isinstance(res["inventory"], list):
                results_inventory.extend(res["inventory"])
            if "issueReports" in res and isinstance(res["issueReports"], list):
                results_issues.extend(res["issueReports"])
                
            print(f"Parsed chunk {i//chunk_size + 1}/{len(combined_raw_text)//chunk_size + 1}: Extracted "
                  f"{len(res.get('suppliers', []))} suppliers, "
                  f"{len(res.get('inventory', []))} inventory items, "
                  f"{len(res.get('issueReports', []))} issues.")
        except Exception as e:
            print(f"Error in chunk {i//chunk_size + 1}: {e}")
            
        time.sleep(3)  # Sleep 3 seconds between chunks to stay strictly below TPM limits
        
    # Write individual files
    with open("extracted_suppliers.json", "w") as f:
        json.dump(results_suppliers, f, indent=2)
    with open("extracted_inventory.json", "w") as f:
        json.dump(results_inventory, f, indent=2)
    with open("extracted_issues.json", "w") as f:
        json.dump(results_issues, f, indent=2)
        
    print("\nExtraction complete!")
    print(f"Suppliers found: {len(results_suppliers)}")
    print(f"Issues found: {len(results_issues)}")
    print(f"Inventory items found: {len(results_inventory)}")

if __name__ == "__main__":
    main()
